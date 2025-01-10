import os
import pandas as pd
from docx import Document
from smtplib import SMTP
from email.message import EmailMessage


def create_documents(template_path, csv_path, output_dir):
    """Generates personalized documents from a template."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Load recipient data
    data = pd.read_csv(csv_path)
    document_template = Document(template_path)

    generated_files = []
    for _, row in data.iterrows():
        # Create a copy of the document template
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            # Replace placeholders with actual data
            for key, value in row.items():
                paragraph.text = paragraph.text.replace(f"{{{{ {key} }}}}", str(value))

        # Save the generated document
        output_file = os.path.join(output_dir, f"{row['name']}_letter.docx")
        doc.save(output_file)
        generated_files.append((row['email'], output_file))

    print(f"Generated {len(generated_files)} documents in {output_dir}")
    return generated_files


def send_emails(generated_files, subject, sender_email, sender_password):
    """Sends emails with the generated documents as attachments."""
    smtp_server = "smtp.gmail.com"
    smtp_port = 587

    with SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(sender_email, sender_password)

        for recipient_email, document_path in generated_files:
            # Extract recipient's name from the document filename
            recipient_name = os.path.basename(document_path).split('_')[0]

            # Create a personalized email body
            email_body = f"Dear {recipient_name},\n\nPlease find your personalized letter attached."

            # Create the email
            msg = EmailMessage()
            msg["From"] = sender_email
            msg["To"] = recipient_email
            msg["Subject"] = subject
            msg.set_content(email_body)

            # Attach the document
            with open(document_path, "rb") as f:
                file_data = f.read()
                file_name = os.path.basename(document_path)
            msg.add_attachment(file_data, maintype="application", subtype="octet-stream", filename=file_name)

            server.send_message(msg)
            print(f"Email sent to {recipient_name} at {recipient_email}")



# Example usage
template = "template.docx"  # Word template with placeholders like {{ name }}, {{ date }}
data_csv = "recipients.csv"  # CSV file with columns matching placeholders
output_directory = "output_letters"
email_subject = "Personalized Letter"
sender = "USERNAME@gmail.com"
password = "PASSWORD"

# Generate documents and send emails
generated_docs = create_documents(template, data_csv, output_directory)
send_emails(generated_docs, email_subject, sender, password)
